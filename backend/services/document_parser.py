"""
Service - Document Parser
==========================
Handles extraction of raw text from uploaded documents (PDF, DOCX, TXT).
Supports metadata extraction and robust error recovery for malformed files.
"""

from __future__ import annotations

import logging
import os
import shutil
import uuid
from pathlib import Path
from typing import Any, Dict

from backend.config import settings
from backend.utils.constants import SUPPORTED_FILE_TYPES
from backend.utils.decorators import log_execution
from backend.utils.exceptions import DocumentProcessingError
from backend.utils.helpers import get_file_extension, sanitize_filename

logger = logging.getLogger(__name__)


@log_execution
def extract_pdf(file_path: str) -> str:
    """Extract text from a PDF file using PyPDF2 with pdfplumber fallback.

    Parameters
    ----------
    file_path : str
        Absolute path to the PDF file.

    Returns
    -------
    str
        Extracted plain text.

    Raises
    ------
    DocumentProcessingError
        If the PDF cannot be read.
    """
    text_parts: list[str] = []

    # Primary: PyPDF2
    try:
        import PyPDF2

        with open(file_path, "rb") as fh:
            reader = PyPDF2.PdfReader(fh)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        if text_parts:
            logger.info("PyPDF2 extracted %d pages from %s", len(text_parts), file_path)
            return "\n\n".join(text_parts)
    except Exception as exc:
        logger.warning("PyPDF2 failed for %s: %s — falling back to pdfplumber", file_path, exc)

    # Fallback: pdfplumber (better table / scanned support)
    try:
        import pdfplumber

        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        if text_parts:
            logger.info("pdfplumber extracted %d pages from %s", len(text_parts), file_path)
            return "\n\n".join(text_parts)
    except Exception as exc:
        logger.error("pdfplumber also failed for %s: %s", file_path, exc)

    if not text_parts:
        raise DocumentProcessingError(
            message="Could not extract text from PDF.",
            detail=f"Both PyPDF2 and pdfplumber failed for {file_path}",
        )
    return "\n\n".join(text_parts)


@log_execution
def extract_docx(file_path: str) -> str:
    """Extract text from a Microsoft Word DOCX file.

    Parameters
    ----------
    file_path : str
        Absolute path to the .docx file.

    Returns
    -------
    str
        Extracted plain text (paragraphs joined by newlines).
    """
    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        text = "\n\n".join(paragraphs)
        logger.info("Extracted %d paragraphs from DOCX %s", len(paragraphs), file_path)
        return text
    except Exception as exc:
        raise DocumentProcessingError(
            message="Failed to extract text from DOCX.",
            detail=str(exc),
        )


@log_execution
def extract_txt(file_path: str) -> str:
    """Read plain text from a .txt file.

    Tries UTF-8 first, then falls back to latin-1.
    """
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            text = Path(file_path).read_text(encoding=encoding)
            logger.info("Read TXT file %s with encoding %s", file_path, encoding)
            return text
        except UnicodeDecodeError:
            continue
    raise DocumentProcessingError(
        message="Failed to read text file.",
        detail=f"Could not decode {file_path} with any supported encoding.",
    )


def extract_text(file_path: str, file_type: str) -> str:
    """Dispatch to the correct extractor based on *file_type*."""
    extractors = {
        "pdf": extract_pdf,
        "docx": extract_docx,
        "txt": extract_txt,
    }
    extractor = extractors.get(file_type)
    if extractor is None:
        raise DocumentProcessingError(
            message=f"Unsupported file type: {file_type}",
            detail=f"Supported types: {SUPPORTED_FILE_TYPES}",
        )
    return extractor(file_path)


def get_document_metadata(file_path: str) -> Dict[str, Any]:
    """Return basic metadata about a file on disk."""
    p = Path(file_path)
    stat = p.stat()
    return {
        "filename": p.name,
        "extension": p.suffix.lstrip(".").lower(),
        "size_bytes": stat.st_size,
        "modified_at": stat.st_mtime,
    }


def save_upload_file(file, filename: str | None = None) -> str:
    """Persist an uploaded ``UploadFile`` (or file-like) to the upload directory.

    Returns the absolute path of the saved file.
    """
    upload_dir = settings.upload_path
    safe_name = sanitize_filename(filename or getattr(file, "filename", "document"))
    unique_name = f"{uuid.uuid4().hex[:8]}_{safe_name}"
    dest = upload_dir / unique_name

    try:
        # FastAPI UploadFile
        if hasattr(file, "file"):
            with open(dest, "wb") as out:
                shutil.copyfileobj(file.file, out)
        else:
            # Raw bytes
            with open(dest, "wb") as out:
                out.write(file.read() if hasattr(file, "read") else file)

        logger.info("Saved upload to %s (%d bytes)", dest, dest.stat().st_size)
        return str(dest)
    except Exception as exc:
        raise DocumentProcessingError(
            message="Failed to save uploaded file.",
            detail=str(exc),
        )
